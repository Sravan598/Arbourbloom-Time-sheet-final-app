"""
Generate FRD (Functional Requirements Document) for AurborBloom
in DOCX format.
"""
import os
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Re-use helpers by importing from sibling script
import importlib.util
spec = importlib.util.spec_from_file_location(
    "brd_helpers", "/app/documents/aurborbloom/scripts/generate_brd_docx.py")
mod = importlib.util.module_from_spec(spec)
spec.loader.exec_module(mod)

PRIMARY = mod.PRIMARY
SECONDARY = mod.SECONDARY
DARK = mod.DARK
GRAY = mod.GRAY
WHITE = mod.WHITE
ASSETS = mod.ASSETS
LOGO = mod.LOGO
OUT_DIR = mod.OUT_DIR

add_heading = mod.add_heading
add_para = mod.add_para
add_bullets = mod.add_bullets
add_numbered = mod.add_numbered
add_table = mod.add_table
add_image = mod.add_image
add_screenshot = mod.add_screenshot
add_portal_appendix = mod.add_portal_appendix
SCREENSHOTS_DIR = mod.SCREENSHOTS_DIR
add_page_break = mod.add_page_break
add_section_callout = mod.add_section_callout
setup_header_footer = mod.setup_header_footer
cover_page = mod.cover_page
shade_cell = mod.shade_cell


def build_frd():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    cover_page(doc, "Functional Requirements\nDocument (FRD)",
               "AB-FRD-2026-001", "1.0")
    setup_header_footer(doc, "Functional Requirements Document")

    # 1. Document Control
    add_heading(doc, "1. Document Control", level=1)
    add_table(doc,
        ["Version", "Date", "Author", "Reviewer", "Description of Change"],
        [
            ["0.1", "2026-01-18", "Lead Engineer", "Project Manager", "Initial draft from BRD baseline"],
            ["0.5", "2026-01-26", "Lead Engineer", "PS VP Engineering", "Tech review"],
            ["1.0", date.today().strftime("%Y-%m-%d"), "Project Manager", "AurborBloom Steering Committee", "Approved baseline"],
        ],
        col_widths=[Inches(0.7), Inches(1.1), Inches(1.7), Inches(1.7), Inches(1.3)]
    )

    add_heading(doc, "1.1 Reference Documents", level=2)
    add_bullets(doc, [
        "AB-BRD-2026-001 — AurborBloom Business Requirements Document, v1.0",
        "AurborBloom HRMS Architecture Specification (Internal)",
        "AurborBloom HRMS Security & Compliance Whitepaper (Internal)",
        "AurborBloom Brand Guidelines (Internal)",
    ])

    add_page_break(doc)

    # 2. Introduction
    add_heading(doc, "2. Introduction", level=1)
    add_heading(doc, "2.1 Purpose", level=2)
    add_para(doc,
        "This Functional Requirements Document (FRD) translates the business requirements captured in the BRD into "
        "discrete, testable functional specifications for the AurborBloom HRMS platform as it will be deployed for "
        "AurborBloom. Each requirement is uniquely identified, mapped back to one or more business "
        "requirements, and accompanied by acceptance criteria that form the basis of UAT.")

    add_heading(doc, "2.2 Intended Audience", level=2)
    add_bullets(doc, [
        "AurborBloom Engineering and Engineering teams (technical reviewers)",
        "AurborBloom People Ops Operations (process validators)",
        "Internal delivery, engineering, and QA teams (implementers)",
        "Internal and external auditors (compliance reviewers)",
    ])

    add_heading(doc, "2.3 Document Conventions", level=2)
    add_bullets(doc, [
        "Functional requirements are prefixed FR-NNN.",
        "Non-functional requirements are prefixed NFR-NNN.",
        "Use cases are prefixed UC-NNN. User stories are prefixed US-NNN.",
        "Priority follows MoSCoW: Must / Should / Could / Won't.",
        "Wherever a requirement traces back to a BRD requirement, the BR-NNN identifier is provided.",
    ])

    add_page_break(doc)

    # 3. System overview
    add_heading(doc, "3. System Overview", level=1)
    add_para(doc,
        "The AurborBloom HRMS is a multi-tenant SaaS platform that comprising a React single-page application, a FastAPI backend, and a "
        "MongoDB persistence tier, all deployed on Kubernetes. For AurborBloom the platform is delivered as a "
        "fully white-labeled tenant identified by the slug 'aurborbloom', accessible at the custom domain "
        "hr.aurborbloom.com. The figure below depicts the overall architecture.")
    add_image(doc, f"{ASSETS}/architecture.png", width_inches=6.4,
              caption="Figure 1 — System Architecture (Multi-Tenant SaaS)")

    add_screenshot(doc, f"{SCREENSHOTS_DIR}/01_tenant_landing.png",
                   "Live Portal — Tenant Landing (hr.aurborbloom.com)",
                   "Figure 1a — Multi-tenant theming rendered live for AurborBloom.")

    add_heading(doc, "3.1 Technology Stack", level=2)
    add_table(doc, ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "React 18, TailwindCSS, react-router-dom, Framer Motion", "User interface, animations, routing"],
            ["State", "React Context API, hooks", "Auth context, tenant context, theming"],
            ["Backend", "FastAPI (Python 3.11), Pydantic v2", "REST API, validation, OpenAPI"],
            ["Database", "MongoDB (test_database)", "Tenant-scoped persistence"],
            ["Auth", "JWT (HS256), bcrypt password hashing", "Stateless authentication"],
            ["AI", "Emergent Universal Key (Claude / GPT / Gemini)", "AI assistant inside HRMS"],
            ["PDF", "ReportLab", "Server-side document generation"],
            ["Calendar", "react-big-calendar", "Calendar UX"],
            ["Hosting", "Kubernetes (Emergent Platform)", "Orchestration, ingress, TLS"],
        ],
        col_widths=[Inches(1.2), Inches(2.6), Inches(2.6)]
    )

    add_heading(doc, "3.2 Tenant Resolution & White-Labeling", level=2)
    add_bullets(doc, [
        "Tenant resolution can occur via slug path (/aurborbloom/...) or custom domain (hr.aurborbloom.com).",
        "On page load the frontend calls GET /api/tenants/by-slug/{slug} or GET /api/tenants/by-domain to fetch tenant metadata.",
        "Tenant metadata drives logo, primary/secondary colors, browser tab title, and feature flags.",
        "All API calls implicitly carry tenant context via the JWT 'tenant_id' claim; cross-tenant access is rejected at the route layer.",
    ])

    add_page_break(doc)

    # 4. User Roles & Permissions
    add_heading(doc, "4. User Roles & Permission Matrix", level=1)
    add_para(doc,
        "Roles below are scoped to the AurborBloom tenant. The Super Admin role exists at the platform level "
        "and is reserved for IT operations.")
    add_table(doc, ["Capability", "Super Admin", "Tenant Admin", "HR Manager", "People Manager", "Employee"],
        [
            ["Manage tenants & branding", "YES", "—", "—", "—", "—"],
            ["Create/Edit Employees", "—", "YES", "YES", "—", "—"],
            ["Approve Timesheet", "—", "YES", "YES", "YES (team)", "—"],
            ["Approve Leave", "—", "YES", "YES", "YES (team)", "—"],
            ["Submit Timesheet / Leave", "—", "YES", "YES", "YES", "YES"],
            ["Raise Tickets", "—", "YES", "YES", "YES", "YES"],
            ["Resolve Tickets", "—", "YES", "YES", "YES (assigned)", "—"],
            ["View All Reports", "YES", "YES", "YES", "Team only", "Self only"],
            ["Configure Departments / Leave Types", "—", "YES", "YES", "—", "—"],
            ["Audit Log Access", "YES", "YES", "Read-only", "—", "—"],
        ],
        col_widths=[Inches(2.0), Inches(0.95), Inches(0.95), Inches(0.95), Inches(0.95), Inches(0.85)]
    )

    add_page_break(doc)

    # 5. Functional Modules
    add_heading(doc, "5. Functional Modules", level=1)

    # ---- 5.1 Auth ----
    add_heading(doc, "5.1 Authentication & Onboarding", level=2)
    add_image(doc, f"{ASSETS}/auth_flow.png", width_inches=6.4,
              caption="Figure 2 — Authentication & Tenant Resolution Flow")
    add_screenshot(doc, f"{SCREENSHOTS_DIR}/02_login.png",
                   "Live Portal — Login",
                   "Figure 2a — Tenant-branded login experience.")
    add_screenshot(doc, f"{SCREENSHOTS_DIR}/03_signup.png",
                   "Live Portal — Signup (Invite-Code Gated)",
                   "Figure 2b — Admin signup requires a valid tenant invite code (FR-001 / FR-002).")
    add_table(doc, ["FR ID", "Functional Requirement", "Priority", "Trace"],
        [
            ["FR-001", "The system shall allow new admins to sign up only when supplied with a valid tenant-specific admin invite code.", "Must", "BR-002"],
            ["FR-002", "The system shall validate the invite code against the tenants collection and reject invalid or expired codes with a 'Invalid admin invite code' message.", "Must", "BR-002"],
            ["FR-003", "The system shall hash passwords using bcrypt with a work factor ≥ 12.", "Must", "BR-019"],
            ["FR-004", "The system shall issue a JWT containing user_id, tenant_id, role, and expiry on successful login.", "Must", "BR-003"],
            ["FR-005", "The system shall lock an account after 5 consecutive failed login attempts within 15 minutes; lock duration is 30 minutes.", "Must", "BR-020"],
            ["FR-006", "The system shall allow password reset through a one-time tokenized link sent to the registered email.", "Should", "BR-019"],
            ["FR-007", "The system shall expire idle sessions after 30 minutes of inactivity.", "Must", "BR-019"],
            ["FR-008", "The system shall display AurborBloom branding (logo, colors, tab title) on all auth pages.", "Must", "BR-001"],
        ],
        col_widths=[Inches(0.7), Inches(4.4), Inches(0.7), Inches(0.6)]
    )

    add_heading(doc, "5.1.1 Use Case — UC-001 Admin Self-Onboarding", level=3)
    add_table(doc, ["Field", "Detail"],
        [
            ["Actor", "Prospective Tenant Admin at AurborBloom"],
            ["Pre-condition", "Valid admin invite code received via secure channel."],
            ["Trigger", "User navigates to hr.aurborbloom.com/signup"],
            ["Main Flow",
             "1. User enters name, email, password, and admin code.\n"
             "2. System validates code against aurborbloom tenant.\n"
             "3. System creates user with role TENANT_ADMIN.\n"
             "4. JWT issued; user redirected to admin dashboard."],
            ["Alternate Flow A1", "Invite code invalid → display error and abort signup."],
            ["Alternate Flow A2", "Email already registered → prompt to log in instead."],
            ["Post-condition", "User account created, audit log entry written."],
            ["Acceptance",
             "Given a valid admin code, when the user submits the signup form with valid data, then the account is created and the dashboard loads within 2 seconds."],
        ],
        col_widths=[Inches(1.5), Inches(5.0)]
    )

    # ---- 5.2 Employee Management ----
    add_heading(doc, "5.2 Employee Management", level=2)
    add_table(doc, ["FR ID", "Functional Requirement", "Priority", "Trace"],
        [
            ["FR-010", "HR shall create an employee record with name, email, department, role, manager, and join date.", "Must", "BR-005"],
            ["FR-011", "HR shall edit any employee field except immutable identifiers.", "Must", "BR-005"],
            ["FR-012", "HR shall deactivate (soft-delete) an employee; deactivated users cannot log in.", "Must", "BR-005"],
            ["FR-013", "HR shall search and filter employees by name, department, manager, status.", "Should", "BR-005"],
            ["FR-014", "Employees shall view and update their own profile photo, contact, and emergency contact.", "Must", "BR-024"],
            ["FR-015", "The system shall enforce uniqueness of email within a tenant.", "Must", "BR-005"],
            ["FR-016", "Department list shall be tenant-configurable; default values for AurborBloom are: General, Engineering, HR, Sales, Marketing, Finance.", "Must", "BR-005"],
        ],
        col_widths=[Inches(0.7), Inches(4.4), Inches(0.7), Inches(0.6)]
    )
    add_screenshot(doc, f"{SCREENSHOTS_DIR}/05_admin_employees.png",
                   "Live Portal — Employee Directory (Admin)",
                   "Figure 5a — Searchable employee directory with role, department, and status filters.")

    # ---- 5.3 Time & Attendance ----
    add_heading(doc, "5.3 Time & Attendance", level=2)
    add_table(doc, ["FR ID", "Functional Requirement", "Priority", "Trace"],
        [
            ["FR-020", "Employees shall log time entries per day with project, task, and hours.", "Must", "BR-006"],
            ["FR-021", "Employees shall save drafts and submit timesheets weekly.", "Must", "BR-006"],
            ["FR-022", "Managers shall approve, reject (with reason), or request changes on a timesheet.", "Must", "BR-007"],
            ["FR-023", "The system shall flag timesheets > 60 hours/week as exceptions.", "Should", "BR-006"],
            ["FR-024", "Approved timesheets shall be locked from edits except by HR.", "Must", "BR-007"],
            ["FR-025", "The system shall produce a weekly utilization report per project.", "Should", "BR-015"],
            ["FR-026", "Time entries shall be exportable in CSV/Excel.", "Must", "BR-015"],
        ],
        col_widths=[Inches(0.7), Inches(4.4), Inches(0.7), Inches(0.6)]
    )
    add_screenshot(doc, f"{SCREENSHOTS_DIR}/14_employee_timesheet.png",
                   "Live Portal — Employee Timesheet (Submitter View)",
                   "Figure 5b — Daily / weekly time entry with project tagging.")
    add_screenshot(doc, f"{SCREENSHOTS_DIR}/06_admin_timesheets.png",
                   "Live Portal — Timesheet Approvals (Manager View)",
                   "Figure 5c — Pending timesheet review queue.")

    # ---- 5.4 Leave ----
    add_heading(doc, "5.4 Leave Management", level=2)
    add_table(doc, ["FR ID", "Functional Requirement", "Priority", "Trace"],
        [
            ["FR-030", "Employees shall view current leave balance per leave type.", "Must", "BR-008"],
            ["FR-031", "Employees shall submit leave requests with type, start, end, and reason.", "Must", "BR-008"],
            ["FR-032", "The system shall reject requests that exceed available balance or overlap existing approved leave.", "Must", "BR-010"],
            ["FR-033", "Managers shall approve or reject requests with optional comments.", "Must", "BR-009"],
            ["FR-034", "Approved leave shall automatically deduct balance and add a calendar entry visible to the team.", "Must", "BR-013"],
            ["FR-035", "The system shall send email + in-app notifications on submit, approval, rejection.", "Must", "BR-017"],
            ["FR-036", "HR shall configure leave types, accrual rules, and carry-over policies per leave type.", "Must", "BR-008"],
            ["FR-037", "Default leave types for AurborBloom: Vacation, Sick, Personal, Unpaid.", "Must", "BR-008"],
        ],
        col_widths=[Inches(0.7), Inches(4.4), Inches(0.7), Inches(0.6)]
    )
    add_image(doc, f"{ASSETS}/leave_flow.png", width_inches=6.2,
              caption="Figure 3 — Leave Workflow (re-shown for module context)")
    add_screenshot(doc, f"{SCREENSHOTS_DIR}/12_employee_leave.png",
                   "Live Portal — Employee Leave (Requester View)",
                   "Figure 5d — Real-time balance and request submission.")
    add_screenshot(doc, f"{SCREENSHOTS_DIR}/07_admin_leave.png",
                   "Live Portal — Leave Approval Queue (Manager View)",
                   "Figure 5e — Manager queue with balance and overlap awareness.")

    # ---- 5.5 Tickets ----
    add_heading(doc, "5.5 Tickets / Helpdesk", level=2)
    add_table(doc, ["FR ID", "Functional Requirement", "Priority", "Trace"],
        [
            ["FR-040", "Employees shall raise a ticket with subject, description, category, priority, and attachments.", "Must", "BR-011"],
            ["FR-041", "Tickets shall be auto-routed by category to the corresponding department's queue.", "Must", "BR-011"],
            ["FR-042", "Agents shall comment, change status, reassign, and close tickets.", "Must", "BR-011"],
            ["FR-043", "Each ticket shall track first-response and resolution SLA, and visually warn when at risk.", "Should", "BR-012"],
            ["FR-044", "Requesters shall receive notifications on status change and may reopen within 7 days of closure.", "Must", "BR-017"],
            ["FR-045", "The system shall support a basic chat-style threaded conversation per ticket (polling-based).", "Must", "BR-011"],
        ],
        col_widths=[Inches(0.7), Inches(4.4), Inches(0.7), Inches(0.6)]
    )
    add_screenshot(doc, f"{SCREENSHOTS_DIR}/15_employee_tickets.png",
                   "Live Portal — Employee Tickets",
                   "Figure 5f — Raise and track support tickets.")
    add_screenshot(doc, f"{SCREENSHOTS_DIR}/08_admin_tickets.png",
                   "Live Portal — Admin / Agent Ticket Queue",
                   "Figure 5g — Admin ticket queue with SLA visibility.")

    # ---- 5.6 Calendar ----
    add_heading(doc, "5.6 Calendar & Scheduling", level=2)
    add_table(doc, ["FR ID", "Functional Requirement", "Priority", "Trace"],
        [
            ["FR-050", "The system shall display a unified calendar combining approved leaves, company events, and personal events.", "Must", "BR-013"],
            ["FR-051", "Users shall create personal and team events with title, start, end, attendees.", "Must", "BR-013"],
            ["FR-052", "Managers shall view team availability heatmap by week.", "Should", "BR-016"],
            ["FR-053", "Calendar shall offer Day, Week, Month, and Agenda views.", "Must", "BR-013"],
        ],
        col_widths=[Inches(0.7), Inches(4.4), Inches(0.7), Inches(0.6)]
    )
    add_screenshot(doc, f"{SCREENSHOTS_DIR}/09_admin_calendar.png",
                   "Live Portal — Unified Calendar",
                   "Figure 5h — Approved leaves, company events, personal events on one calendar.")

    # ---- 5.7 Documents ----
    add_heading(doc, "5.7 Document Vault", level=2)
    add_table(doc, ["FR ID", "Functional Requirement", "Priority", "Trace"],
        [
            ["FR-060", "Employees shall download personal documents such as offer letters, payslips, policy acknowledgments.", "Must", "BR-014"],
            ["FR-061", "HR shall upload company-wide policy documents available to all employees.", "Must", "BR-014"],
            ["FR-062", "All documents shall be stored encrypted at rest and accessible only to authorized roles.", "Must", "BR-019"],
            ["FR-063", "Document downloads shall be audit-logged.", "Must", "BR-018"],
        ],
        col_widths=[Inches(0.7), Inches(4.4), Inches(0.7), Inches(0.6)]
    )
    add_screenshot(doc, f"{SCREENSHOTS_DIR}/13_employee_documents.png",
                   "Live Portal — Document Vault",
                   "Figure 5i — PIN-secured personal document access.")

    # ---- 5.8 Performance ----
    add_heading(doc, "5.8 Performance Reviews", level=2)
    add_table(doc, ["FR ID", "Functional Requirement", "Priority", "Trace"],
        [
            ["FR-070", "Managers shall create review cycles with self, peer, and manager assessments.", "Should", "BG-07"],
            ["FR-071", "Employees shall set quarterly goals (OKRs) and update progress.", "Should", "BG-07"],
            ["FR-072", "HR shall view aggregate performance trends by department.", "Could", "BG-10"],
        ],
        col_widths=[Inches(0.7), Inches(4.4), Inches(0.7), Inches(0.6)]
    )

    # ---- 5.9 Reports ----
    add_heading(doc, "5.9 Reports & Analytics", level=2)
    add_table(doc, ["FR ID", "Functional Requirement", "Priority", "Trace"],
        [
            ["FR-080", "The system shall provide standard reports: Headcount, Attendance, Leave Utilization, Ticket SLA, Project Utilization.", "Must", "BR-015"],
            ["FR-081", "Reports shall be filterable by date range, department, and project.", "Must", "BR-015"],
            ["FR-082", "Reports shall be exportable as CSV and Excel.", "Must", "BR-015"],
            ["FR-083", "Dashboards shall refresh in near real-time (< 1 minute lag).", "Should", "BR-016"],
        ],
        col_widths=[Inches(0.7), Inches(4.4), Inches(0.7), Inches(0.6)]
    )

    add_page_break(doc)

    # 6. User Stories
    add_heading(doc, "6. User Stories with Acceptance Criteria", level=1)
    stories = [
        ("US-01",
         "As an Employee, I want to submit a leave request from my phone so that I do not need to be at my desk.",
         "Given I am authenticated on a mobile browser, when I open Leave → New Request, then I can submit a request and receive a confirmation toast within 2 seconds."),
        ("US-02",
         "As a Manager, I want a single approval queue so that I do not switch contexts between leave, time, and tickets.",
         "Given I have pending approvals across modules, when I open My Approvals, then I see a unified list ordered by SLA risk."),
        ("US-03",
         "As HR, I want to bulk import employees from a spreadsheet so that I avoid manual entry during onboarding waves.",
         "Given a CSV in the standard template, when I upload it, then validated rows are imported and errors are downloadable as a per-row report."),
        ("US-04",
         "As an Admin, I want browser tab titles to display 'AurborBloom — HR Portal' so that brand identity is preserved.",
         "Given any page on hr.aurborbloom.com, then the tab title contains 'AurborBloom' and never a generic platform string."),
        ("US-05",
         "As an Auditor, I want to download an immutable audit log for any date range so that I can reconcile changes.",
         "Given Auditor role, when I request the export, then a CSV with timestamp, actor, action, before/after states is produced."),
        ("US-06",
         "As an Employee, I want to see my real-time leave balance before I submit a request so that I avoid rejections.",
         "Given my leave page, when I select a leave type and dates, then the projected remaining balance is shown live."),
        ("US-07",
         "As a Manager, I want to be notified within 1 minute of an SLA-at-risk ticket on my queue.",
         "Given a ticket within 80% of its SLA window, when polling fires, then the agent receives an in-app warning and an email."),
        ("US-08",
         "As HR, I want to configure leave types per policy so that I can adapt to state-specific sick-leave laws.",
         "Given Tenant Admin role, when I edit leave types, then changes apply prospectively without affecting historical balances."),
    ]
    add_table(doc, ["ID", "User Story", "Acceptance Criteria"],
        [list(s) for s in stories],
        col_widths=[Inches(0.7), Inches(2.6), Inches(3.0)]
    )

    add_page_break(doc)

    # 7. Data Model
    add_heading(doc, "7. Data Model", level=1)
    add_image(doc, f"{ASSETS}/erd.png", width_inches=6.4,
              caption="Figure 4 — Core Data Model (Simplified ER Diagram)")
    add_para(doc,
        "All collections are tenant-scoped: every persisted document carries an indexed tenant_id field. "
        "Backend queries automatically constrain by the JWT-resolved tenant_id; integration tests verify "
        "that no API call may return another tenant's data.")

    add_heading(doc, "7.1 Key Collections", level=2)
    add_table(doc, ["Collection", "Cardinality", "Indexes", "Notes"],
        [
            ["tenants", "1 per company", "slug (unique), custom_domain (unique sparse)", "Stores branding, settings, invite code"],
            ["users", "1..N per tenant", "tenant_id+email (unique)", "Auth + role"],
            ["employees", "1..1 with users", "tenant_id, manager_id", "Profile data"],
            ["timesheets", "Many", "tenant_id+employee_id+date", "Locked after approval"],
            ["leave_requests", "Many", "tenant_id+employee_id, status", "Approval workflow"],
            ["tickets", "Many", "tenant_id+status, assigned_to", "SLA-tracked"],
            ["calendar_events", "Many", "tenant_id+start", "Includes leave-derived events"],
            ["documents", "Many", "tenant_id+owner_id", "Encrypted at rest"],
            ["audit_logs", "Many (append-only)", "tenant_id+timestamp", "Immutable"],
        ],
        col_widths=[Inches(1.2), Inches(1.2), Inches(2.2), Inches(2.0)]
    )

    add_page_break(doc)

    # 8. APIs
    add_heading(doc, "8. API Surface (Representative)", level=1)
    add_para(doc,
        "All backend endpoints are prefixed with /api and require a valid JWT (except the explicitly public ones). "
        "Tenant context is derived from the JWT's tenant_id claim. The list below is representative, not exhaustive; "
        "the full OpenAPI specification is generated by FastAPI at /api/docs.")
    add_table(doc, ["Method", "Endpoint", "Purpose", "Auth"],
        [
            ["POST", "/api/auth/signup", "Tenant admin sign-up with invite code", "Public"],
            ["POST", "/api/auth/login", "Login, returns JWT", "Public"],
            ["GET",  "/api/tenants/by-slug/{slug}", "Resolve tenant by slug", "Public"],
            ["GET",  "/api/tenants/by-domain", "Resolve tenant by custom domain", "Public"],
            ["GET",  "/api/employees", "List tenant employees", "Authenticated"],
            ["POST", "/api/employees", "Create employee", "Tenant Admin / HR"],
            ["GET",  "/api/timesheets", "List timesheets (filtered)", "Authenticated"],
            ["POST", "/api/timesheets", "Create / submit timesheet", "Employee+"],
            ["POST", "/api/timesheets/{id}/approve", "Approve timesheet", "Manager+"],
            ["GET",  "/api/leave/balance", "Get leave balance", "Authenticated"],
            ["POST", "/api/leave/request", "Submit leave request", "Employee+"],
            ["POST", "/api/leave/{id}/approve", "Approve leave", "Manager+"],
            ["POST", "/api/tickets", "Create ticket", "Authenticated"],
            ["GET",  "/api/tickets/{id}", "Get ticket detail incl. comments", "Authenticated"],
            ["GET",  "/api/calendar/events", "List events for visible range", "Authenticated"],
            ["GET",  "/api/reports/{type}", "Generate report (CSV/JSON)", "HR / Admin"],
            ["GET",  "/api/audit-logs", "Read-only audit log", "Admin / Auditor"],
        ],
        col_widths=[Inches(0.7), Inches(2.6), Inches(2.4), Inches(0.9)]
    )

    add_page_break(doc)

    # 9. Non-functional
    add_heading(doc, "9. Non-Functional Requirements", level=1)
    add_table(doc, ["NFR ID", "Category", "Requirement", "Target"],
        [
            ["NFR-001", "Performance", "API response under normal load", "p95 < 500 ms; p99 < 1.2 s"],
            ["NFR-002", "Performance", "Time-to-interactive (TTI) on dashboard", "< 2.5 s on 4G"],
            ["NFR-003", "Availability", "Monthly uptime (excluding planned maintenance)", "≥ 99.9%"],
            ["NFR-004", "Scalability", "Concurrent active users supported per tenant", "≥ 1,000"],
            ["NFR-005", "Scalability", "Records per collection without performance degradation", "≥ 10 million"],
            ["NFR-006", "Security", "Password hashing", "bcrypt (cost ≥ 12)"],
            ["NFR-007", "Security", "Transport encryption", "TLS 1.2+ enforced"],
            ["NFR-008", "Security", "Storage encryption", "AES-256 at rest for PII fields"],
            ["NFR-009", "Security", "Tenant isolation", "100% — verified by automated tests"],
            ["NFR-010", "Privacy", "Data subject access response time", "≤ 30 days"],
            ["NFR-011", "Reliability", "RPO (Recovery Point Objective)", "≤ 1 hour"],
            ["NFR-012", "Reliability", "RTO (Recovery Time Objective)", "≤ 4 hours"],
            ["NFR-013", "Usability", "Accessibility", "WCAG 2.1 AA"],
            ["NFR-014", "Compatibility", "Supported browsers", "Latest 2 versions of Chrome, Edge, Firefox, Safari"],
            ["NFR-015", "Mobile", "Responsive layout breakpoints", "≥ 360 px width"],
            ["NFR-016", "Audit", "Audit log retention", "≥ 7 years"],
            ["NFR-017", "Maintainability", "Backend test coverage", "≥ 80% for routes & services"],
            ["NFR-018", "Observability", "Structured logging + metrics + traces", "100% of API endpoints"],
            ["NFR-019", "Localization", "Default time zone", "America/New_York with override"],
            ["NFR-020", "Compliance", "Audit-ready evidence collection", "Available within 5 business days"],
        ],
        col_widths=[Inches(0.8), Inches(1.2), Inches(2.8), Inches(1.7)]
    )

    add_page_break(doc)

    # 10. Risks
    add_heading(doc, "10. Functional Risk Analysis", level=1)
    add_table(doc, ["ID", "Risk", "Probability", "Impact", "Mitigation"],
        [
            ["FRK-01", "Race condition on leave balance during concurrent submissions",
             "Medium", "High",
             "Optimistic concurrency check on balance, server-side authoritative deduction in single transaction."],
            ["FRK-02", "Browser tab title regressions when tenant context changes",
             "Medium", "Low",
             "Centralized useBrowserTitle hook keyed on tenant resolution."],
            ["FRK-03", "Custom domain CNAME misconfiguration",
             "Medium", "Medium",
             "Pre-flight DNS check + admin-visible verification status."],
            ["FRK-04", "Incomplete role enforcement on new endpoints",
             "Medium", "High",
             "RBAC decorator pattern enforced via lint rule + integration tests."],
            ["FRK-05", "MongoDB ObjectId leakage in API responses",
             "Low", "Medium",
             "Pydantic response models, _id excluded by default, contract tests."],
            ["FRK-06", "Duplicate /auth/signup routes leading to inconsistent validation",
             "Confirmed", "Medium",
             "Refactor server.py to remove duplicate; consolidate into routes/auth.py."],
            ["FRK-07", "Email deliverability for notifications",
             "Medium", "Medium",
             "Use authenticated SMTP/Resend with SPF, DKIM, DMARC."],
        ],
        col_widths=[Inches(0.7), Inches(2.0), Inches(0.9), Inches(0.7), Inches(2.4)]
    )

    add_page_break(doc)

    # 11. Traceability
    add_heading(doc, "11. Traceability Matrix (BR → FR → UC / US)", level=1)
    add_para(doc,
        "Every business requirement from the BRD is satisfied by one or more functional requirements, validated by "
        "use cases or user stories, and ultimately verified by acceptance test cases.")
    add_table(doc, ["Business Req.", "Functional Req.", "Use Cases / Stories", "Test Method"],
        [
            ["BR-001 Branding", "FR-008", "UC-001, US-04", "Visual + automated regression"],
            ["BR-002 Invite-only signup", "FR-001, FR-002", "UC-001", "Functional + negative tests"],
            ["BR-003 Auth required", "FR-004, FR-007", "UC-001", "API + UI session tests"],
            ["BR-004 RBAC", "FR-001..FR-008, role matrix", "UC-001", "Role permutation tests"],
            ["BR-005 Employee CRUD", "FR-010..FR-016", "US-03", "CRUD test suite"],
            ["BR-006/7 Time", "FR-020..FR-026", "US-02", "Workflow + approval tests"],
            ["BR-008..10 Leave", "FR-030..FR-037", "US-01, US-06, US-08", "Workflow + balance tests"],
            ["BR-011/12 Tickets", "FR-040..FR-045", "US-07", "SLA + workflow tests"],
            ["BR-013 Calendar", "FR-050..FR-053", "US-01", "Cross-module integration"],
            ["BR-014 Documents", "FR-060..FR-063", "—", "Permission + audit tests"],
            ["BR-015/16 Reports", "FR-080..FR-083", "—", "Snapshot + export tests"],
            ["BR-017 Notifications", "FR-035, FR-044", "US-07", "Email + in-app tests"],
            ["BR-018 Audit", "FR-063 + audit_logs", "US-05", "Append-only verification"],
            ["BR-019/20 Security", "FR-003..FR-005, NFR-006..009", "—", "Pen test + automated checks"],
            ["BR-021 Availability", "NFR-003", "—", "Synthetic monitoring"],
            ["BR-022 Performance", "NFR-001, NFR-002", "—", "Load test"],
            ["BR-023 Accessibility", "NFR-013", "—", "Axe / Lighthouse audit"],
            ["BR-024 Self-service", "FR-014, FR-031, FR-040", "US-01", "End-to-end flow tests"],
            ["BR-025 Integrations", "API surface", "—", "Contract tests"],
            ["BR-026 Privacy", "NFR-010", "US-05", "DSAR drill"],
            ["BR-027 Localization", "NFR-019", "—", "TZ regression tests"],
            ["BR-028 Mobile", "NFR-015", "US-01", "Mobile viewport tests"],
            ["BR-029 Tenant isolation", "NFR-009", "—", "Cross-tenant fuzz tests"],
            ["BR-030 DR", "NFR-011, NFR-012", "—", "DR drill"],
        ],
        col_widths=[Inches(1.7), Inches(1.6), Inches(1.6), Inches(1.5)]
    )

    add_page_break(doc)

    # 12. UAT
    add_heading(doc, "12. UAT Strategy & Exit Criteria", level=1)
    add_bullets(doc, [
        "Three UAT cycles: smoke, full regression, and pre-prod soak.",
        "Test data set includes 50 sample employees, 5 managers, 3 departments to start.",
        "Defect severity classes: S1 Blocker, S2 Major, S3 Minor, S4 Cosmetic.",
        "Exit criteria: Zero S1, zero S2, < 5 S3 open at sign-off.",
        "Performance gate: NFR-001/002 met under 2× expected load.",
        "Security gate: Pen test report with zero High findings outstanding.",
    ])

    # 13. Glossary
    add_heading(doc, "13. Glossary", level=1)
    add_table(doc, ["Term", "Definition"],
        [
            ["FR / NFR", "Functional / Non-Functional Requirement"],
            ["JWT", "JSON Web Token — signed, stateless auth token"],
            ["MoSCoW", "Must / Should / Could / Won't priority framework"],
            ["RBAC", "Role-Based Access Control"],
            ["SLA", "Service Level Agreement"],
            ["UAT", "User Acceptance Testing"],
            ["WCAG", "Web Content Accessibility Guidelines"],
            ["RPO/RTO", "Recovery Point / Time Objectives"],
            ["DSAR", "Data Subject Access Request"],
        ],
        col_widths=[Inches(1.4), Inches(5.0)]
    )

    # 14. Sign-off
    add_page_break(doc)
    add_heading(doc, "14. Approval & Sign-off", level=1)
    add_para(doc,
        "By signing below, the parties acknowledge that this Functional Requirements Document accurately translates "
        "the business needs into implementable functional and non-functional specifications.")
    add_table(doc, ["Role", "Name", "Signature", "Date"],
        [
            ["VP Engineering", "AurborBloom", "_____________________", "__________"],
            ["Chief Information Security Officer", "AurborBloom", "_____________________", "__________"],
            ["Head of People Operations", "AurborBloom", "_____________________", "__________"],
            ["Lead Engineer", "Internal Engineering", "_____________________", "__________"],
            ["Product Manager", "Product Office", "_____________________", "__________"],
        ],
        col_widths=[Inches(1.7), Inches(2.4), Inches(1.6), Inches(0.8)]
    )

    out = f"{OUT_DIR}/AurborBloom_FRD_v1.0.docx"
    doc.save(out)
    print(f"FRD DOCX saved → {out}")
    return out


if __name__ == "__main__":
    build_frd()
