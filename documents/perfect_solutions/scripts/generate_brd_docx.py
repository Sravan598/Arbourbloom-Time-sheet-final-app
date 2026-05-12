"""
Generate BRD (Business Requirements Document) for Perfect Solutions
in both DOCX and PDF formats.
"""
import os
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = "/app/documents/perfect_solutions"
ASSETS = f"{OUT_DIR}/assets"
LOGO = f"{ASSETS}/logo.png"

PRIMARY = RGBColor(0x1E, 0x3A, 0x8A)
SECONDARY = RGBColor(0x3B, 0x82, 0xF6)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x64, 0x74, 0x8B)
LIGHT_GRAY = RGBColor(0xE2, 0xE8, 0xF0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_border(cell, color="1E3A8A", sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc, text, level=1, color=PRIMARY, size=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    else:
        run.font.size = Pt(20 - level * 2)
    run.font.name = 'Calibri'
    if level == 1:
        # add bottom border line
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:color'), '1E3A8A')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_para(doc, text, bold=False, italic=False, size=11, color=DARK, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        run = p.runs[0] if p.runs else p.add_run(it)
        run.text = it
        run.font.size = Pt(11)
        run.font.color.rgb = DARK


def add_numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        if p.runs:
            p.runs[0].text = it
        else:
            p.add_run(it)
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.color.rgb = DARK


def add_table(doc, headers, rows, col_widths=None, header_fill="1E3A8A"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        shade_cell(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
            run.font.name = 'Calibri'
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 0:
                shade_cell(cell, "F1F5F9")
    return table


def add_image(doc, path, width_inches=6.4, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_inches))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = c.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9.5)
        cr.font.color.rgb = GRAY


def add_screenshot(doc, path, label, caption, width_inches=6.2):
    """Embed a portal screenshot with a colored label bar and caption."""
    # Label bar (top, dark blue)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(width_inches)
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = WHITE
    r.font.name = 'Calibri'
    shade_cell(cell, "1E3A8A")
    set_cell_border(cell, color="1E3A8A", sz="4")
    # Image
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    p2.add_run().add_picture(path, width=Inches(width_inches))
    # Caption
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_before = Pt(0)
    cr = c.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9.5)
    cr.font.color.rgb = GRAY


SCREENSHOTS_DIR = f"{ASSETS}/screenshots"

PORTAL_SHOTS = [
    ("01_tenant_landing.png", "Screen A1 — Tenant Landing",
     "Branded landing page for hr.perfectsolutions.com — logo, brand colors, and tagline applied automatically by the multi-tenant theming engine."),
    ("02_login.png", "Screen A2 — Login",
     "Login experience inherits the tenant brand. Tenant slug captured implicitly; only Perfect Solutions credentials are accepted."),
    ("03_signup.png", "Screen A3 — Admin Signup (Invite-Code Gated)",
     "New tenant admins must present a tenant-specific invite code to register, enforcing the secure onboarding flow."),
    ("04_admin_dashboard.png", "Screen A4 — Admin Dashboard",
     "Unified admin landing — KPIs, recent activity, and quick access to operations modules."),
    ("05_admin_employees.png", "Screen A5 — Employee Directory (Admin)",
     "Searchable, filterable directory with role and department context. Drives every people-related workflow."),
    ("06_admin_timesheets.png", "Screen A6 — Timesheet Approvals (Admin / Manager)",
     "Pending timesheet review queue. Approve, reject, or request changes; bulk actions supported."),
    ("07_admin_leave.png", "Screen A7 — Leave Approval Queue (Admin / Manager)",
     "Pending leave requests with balance, overlap, and policy context for one-click decisions."),
    ("08_admin_tickets.png", "Screen A8 — Tickets / Helpdesk (Admin View)",
     "Internal helpdesk with categorization, SLA visibility, and threaded responses."),
    ("09_admin_calendar.png", "Screen A9 — Unified Calendar",
     "Approved leaves, company events, and scheduled meetings on a single calendar — Day, Week, Month, Agenda views."),
    ("10_admin_projects.png", "Screen A10 — Projects & Allocation",
     "Project list with capacity and allocation overview — drives the utilization reports."),
    ("11_employee_dashboard.png", "Screen A11 — Employee Dashboard",
     "Self-service home: clock in/out, pending requests, announcements, quick actions."),
    ("12_employee_leave.png", "Screen A12 — Employee Leave",
     "Real-time leave balance with request submission, status tracking, and PDF export."),
    ("13_employee_documents.png", "Screen A13 — Document Vault",
     "PIN-secured personal document vault — payslips, offer letters, signed policies."),
    ("14_employee_timesheet.png", "Screen A14 — Employee Timesheet",
     "Daily/weekly entry with project tagging, drafts, and submit-for-approval workflow."),
    ("15_employee_tickets.png", "Screen A15 — Employee Tickets",
     "Raise and track internal support tickets; threaded comments and status updates."),
]


def add_portal_appendix(doc):
    """Append the full portal walkthrough section."""
    add_page_break(doc)
    add_heading(doc, "Appendix A — Portal Walkthrough (Screenshots)", level=1)
    add_para(doc,
        "This appendix presents annotated screen captures from the live Perfect Solutions tenant of the Perfect Solutions "
        "HRMS portal. Screens are grouped by role to illustrate how the platform satisfies the requirements stated "
        "earlier in this document. All screens are rendered on the perfectsolutions tenant with full brand styling "
        "(logo, primary color #1E3A8A) applied automatically.")

    add_heading(doc, "A.1 Public / Tenant-Branded Pages", level=2)
    for fn, label, cap in PORTAL_SHOTS[0:3]:
        add_screenshot(doc, f"{SCREENSHOTS_DIR}/{fn}", label, cap)
        add_para(doc, "")

    add_page_break(doc)
    add_heading(doc, "A.2 Admin / Manager Views", level=2)
    for fn, label, cap in PORTAL_SHOTS[3:10]:
        add_screenshot(doc, f"{SCREENSHOTS_DIR}/{fn}", label, cap)
        add_para(doc, "")

    add_page_break(doc)
    add_heading(doc, "A.3 Employee Self-Service Views", level=2)
    for fn, label, cap in PORTAL_SHOTS[10:]:
        add_screenshot(doc, f"{SCREENSHOTS_DIR}/{fn}", label, cap)
        add_para(doc, "")


def add_page_break(doc):
    doc.add_page_break()


def add_section_callout(doc, text, fill="DBEAFE"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = PRIMARY
    shade_cell(cell, fill)
    set_cell_border(cell, color="1E3A8A", sz="4")


def setup_header_footer(doc, title):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    header = section.header
    htable = header.add_table(rows=1, cols=2, width=Inches(6.5))
    htable.autofit = True
    left, right = htable.rows[0].cells
    # logo
    left.paragraphs[0].add_run().add_picture(LOGO, width=Inches(1.1))
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rrun = rp.add_run(f"{title}\nPerfect Solutions • Confidential")
    rrun.font.size = Pt(9)
    rrun.font.color.rgb = PRIMARY
    rrun.bold = True

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Perfect Solutions HRMS  |  © 2026  |  Page ")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY
    # page number field
    fr2 = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    fr2._r.append(fldChar1)
    fr2._r.append(instrText)
    fr2._r.append(fldChar2)
    fr2.font.size = Pt(8)
    fr2.font.color.rgb = GRAY


def cover_page(doc, doc_type, doc_code, version):
    # Big colored header band using table
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.add_run().add_picture(LOGO, width=Inches(2.2))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(30)
    r = p2.add_run("PERFECT SOLUTIONS")
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = PRIMARY

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("Unleashing IT Talent")
    r.italic = True; r.font.size = Pt(13); r.font.color.rgb = SECONDARY

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(60)
    r = p4.add_run(doc_type)
    r.bold = True; r.font.size = Pt(36); r.font.color.rgb = DARK

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p5.add_run("Internal HRMS Implementation")
    r.font.size = Pt(16); r.font.color.rgb = GRAY

    # info table
    doc.add_paragraph().paragraph_format.space_before = Pt(60)
    info = doc.add_table(rows=5, cols=2)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_data = [
        ("Document Code", doc_code),
        ("Version", version),
        ("Date", date.today().strftime("%B %d, %Y")),
        ("Classification", "Confidential — Internal Use"),
        ("Prepared By", "IT & Product Implementation Team"),
    ]
    for i, (k, v) in enumerate(info_data):
        c1 = info.rows[i].cells[0]
        c2 = info.rows[i].cells[1]
        c1.text = ''; c2.text = ''
        r1 = c1.paragraphs[0].add_run(k)
        r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = WHITE
        shade_cell(c1, "1E3A8A")
        r2 = c2.paragraphs[0].add_run(v)
        r2.font.size = Pt(11); r2.font.color.rgb = DARK
        shade_cell(c2, "F8FAFC")
    add_page_break(doc)


# =====================================================================
# BRD CONTENT
# =====================================================================
def build_brd():
    doc = Document()
    # default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    cover_page(doc, "Business Requirements\nDocument (BRD)",
               "PS-BRD-2026-001", "1.0")
    setup_header_footer(doc, "Business Requirements Document")

    # 1. Document Control
    add_heading(doc, "1. Document Control", level=1)
    add_table(doc,
        ["Version", "Date", "Author", "Reviewer", "Description of Change"],
        [
            ["0.1", "2026-01-15", "Project Manager", "Perfect Solutions HR", "Initial draft"],
            ["0.5", "2026-01-22", "Implementation Lead", "PS IT Director", "Internal review"],
            ["1.0", date.today().strftime("%Y-%m-%d"), "Project Manager", "PS Steering Committee", "Approved baseline"],
        ],
        col_widths=[Inches(0.7), Inches(1.1), Inches(1.5), Inches(1.5), Inches(1.7)]
    )
    add_para(doc, "")

    add_heading(doc, "1.1 Distribution List", level=2)
    add_table(doc, ["Name / Role", "Organization", "Purpose"],
        [
            ["Director, IT Services", "Perfect Solutions", "Approver"],
            ["Head of Human Resources", "Perfect Solutions", "Approver"],
            ["Chief Financial Officer", "Perfect Solutions", "Reviewer"],
            ["Engineering Manager", "Perfect Solutions", "Reviewer"],
            ["Implementation Lead", "Perfect Solutions IT", "Owner"],
            ["Product Manager", "Perfect Solutions IT", "Author"],
        ],
        col_widths=[Inches(2.2), Inches(2.0), Inches(2.0)]
    )

    add_page_break(doc)

    # 2. Executive Summary
    add_heading(doc, "2. Executive Summary", level=1)
    add_para(doc,
        "Perfect Solutions is a mid-sized IT Services and Software Development firm headquartered in the United States, "
        "operating with a workforce of approximately 51–250 employees distributed across engineering, sales, marketing, "
        "finance, human resources, and general operations. As the company scales its delivery footprint and onboards "
        "increasingly distributed talent, it has identified the need to consolidate its people-operations stack onto a "
        "single, multi-tenant SaaS platform — Perfect Solutions HRMS.")
    add_para(doc,
        "This Business Requirements Document (BRD) captures the consolidated set of business goals, stakeholder "
        "expectations, success metrics, scope boundaries, and business-level requirements that must be satisfied by the "
        "Perfect Solutions HRMS deployment. It serves as the contractual foundation upon which the "
        "Functional Requirements Document (FRD), system design, and acceptance test cases will be built.")
    add_section_callout(doc,
        "Key Outcome Targets — reduce HR administrative effort by 60%, eliminate spreadsheet-driven tracking, "
        "achieve 99.9% platform availability, and deliver a measurable improvement in employee experience within "
        "two quarters of go-live.")

    add_heading(doc, "2.1 Business Drivers", level=2)
    add_bullets(doc, [
        "Rapid headcount growth has outstripped existing spreadsheet- and email-based HR processes.",
        "Multiple disconnected tools (separate apps for leave, attendance, helpdesk, calendar) cause data drift and reconciliation overhead.",
        "Compliance and audit readiness demand a centralized, immutable record of HR transactions.",
        "Distributed teams require a single, branded portal accessible from any device.",
        "Leadership requires real-time, role-based dashboards for workforce decisions.",
    ])

    add_page_break(doc)

    # 3. Company Overview
    add_heading(doc, "3. Company Overview — Perfect Solutions", level=1)
    add_table(doc, ["Attribute", "Value"],
        [
            ["Legal Name", "Perfect Solutions"],
            ["Industry", "IT Services & Software Development"],
            ["Company Size", "Medium (51 – 250 employees)"],
            ["Primary Markets", "United States, with offshore engineering presence"],
            ["Brand Tagline", "Unleashing IT Talent"],
            ["Primary Brand Color", "#1E3A8A (Deep Corporate Blue)"],
            ["Secondary Brand Color", "#3B82F6 (Action Blue)"],
            ["Public Website", "www.perfectgroupus.com"],
            ["HRMS Tenant Slug", "perfectsolutions"],
            ["Custom Domain", "hr.perfectsolutions.com"],
            ["Departments In Scope",
             "Engineering, Sales, Marketing, Finance, Human Resources, General"],
        ],
        col_widths=[Inches(2.0), Inches(4.4)]
    )

    add_heading(doc, "3.1 Strategic Context", level=2)
    add_para(doc,
        "Perfect Solutions positions itself as a high-velocity IT talent firm. The HR function is therefore not a back-office "
        "cost center but a competitive differentiator: speed of hiring, accuracy of billing for time-and-materials engagements, "
        "and employee retention all flow directly through HR systems. The HRMS is therefore treated as a Tier-1 business "
        "system with stringent uptime, security, and reporting requirements.")

    add_screenshot(doc, f"{SCREENSHOTS_DIR}/01_tenant_landing.png",
                   "Live Capture — Perfect Solutions Tenant Landing (hr.perfectsolutions.com)",
                   "Figure 6 — The tenant-branded landing page rendered from live multi-tenant infrastructure.")

    add_image(doc, f"{ASSETS}/modules.png", width_inches=6.4,
              caption="Figure 5 — HRMS Module Map for Perfect Solutions")

    add_page_break(doc)

    # 4. Stakeholders
    add_heading(doc, "4. Stakeholders & User Personas", level=1)
    add_para(doc,
        "The following stakeholder map identifies the primary actors who will interact with, sponsor, govern, or be "
        "impacted by the Perfect Solutions HRMS deployment.")
    add_table(doc, ["Stakeholder", "Role", "Influence", "Interest", "Engagement Strategy"],
        [
            ["Director, IT Services", "Executive Sponsor", "High", "High", "Steering committee, monthly review"],
            ["Head of Human Resources", "Business Owner", "High", "High", "Weekly working group"],
            ["CFO", "Financial Approver", "High", "Medium", "Budget gates, ROI checkpoints"],
            ["Engineering Manager", "Power User", "Medium", "High", "Pilot user, feedback loop"],
            ["HR Operations Lead", "Day-to-day Owner", "Medium", "High", "UAT, training champion"],
            ["IT Security Officer", "Risk Approver", "High", "Medium", "Security review at each phase"],
            ["Employees (all)", "End Users", "Low", "High", "Onboarding sessions, in-app guides"],
            ["Managers (all departments)", "Approvers", "Medium", "High", "Train-the-trainer model"],
            ["External Auditor", "Compliance Reviewer", "Medium", "Low", "Read-only access at audit time"],
        ],
        col_widths=[Inches(1.6), Inches(1.3), Inches(0.8), Inches(0.8), Inches(2.0)]
    )

    add_heading(doc, "4.1 Detailed User Personas", level=2)
    personas = [
        ("Priya — HR Operations Lead",
         "Manages onboarding, leave approvals, policy enforcement, audit responses.",
         "Manual data entry, chasing managers for approvals, reconciling spreadsheets, generating monthly reports.",
         "A single dashboard with pending approvals, automated reminders, and one-click report exports."),
        ("Daniel — Engineering Manager",
         "Approves time, leave, and ticket escalations for an engineering team of ~25.",
         "Switching between tools, no real-time visibility into team capacity, approval requests buried in email.",
         "Mobile-friendly approval queue, team capacity heatmap, project allocation visibility."),
        ("Anita — Software Engineer",
         "Logs hours, requests leave, raises IT helpdesk tickets, downloads payslips.",
         "Forgetting to log time, unclear leave balance, no visibility into ticket status.",
         "Friction-free time entry, instant leave balance, ticket status notifications."),
        ("Marcus — CFO",
         "Approves payroll, monitors workforce cost, reviews compliance posture.",
         "Late or inconsistent reports, lack of trend visibility, manual reconciliation with payroll vendor.",
         "Executive dashboard with cost-per-department, trend charts, exportable payroll registers."),
        ("Sophia — IT Security Officer",
         "Owns access governance, data classification, audit response.",
         "No central audit log, ambiguous role definitions, password sprawl.",
         "Immutable audit trail, granular RBAC, SSO/2FA, encryption at rest and in transit."),
    ]
    for name, role, pain, want in personas:
        add_para(doc, name, bold=True, color=PRIMARY, size=12)
        add_para(doc, f"Role: {role}")
        add_para(doc, f"Pain Points: {pain}")
        add_para(doc, f"Desired Outcome: {want}")
        add_para(doc, "")

    add_page_break(doc)

    # 5. Business Objectives
    add_heading(doc, "5. Business Goals & Success Metrics", level=1)
    add_para(doc,
        "The following SMART business objectives have been agreed by the Perfect Solutions steering committee. Each "
        "objective is paired with a measurable Key Performance Indicator (KPI) and a target value to be achieved within "
        "two quarters of go-live (T+6 months).")

    add_table(doc, ["#", "Business Goal", "KPI", "Baseline", "Target (T+6m)"],
        [
            ["BG-01", "Centralize all HR data into a single platform",
             "% of HR processes digitized in HRMS", "0%", "≥ 95%"],
            ["BG-02", "Automate attendance and timesheet capture",
             "Manual time-entry corrections / month", "~120", "< 15"],
            ["BG-03", "Reduce leave processing cycle time",
             "Avg. days from request to decision", "3.2 days", "< 1 day"],
            ["BG-04", "Improve internal support responsiveness",
             "Avg. ticket first-response time", "8 hours", "< 2 hours"],
            ["BG-05", "Increase employee self-service adoption",
             "% of leave/timesheet actions performed by employees themselves",
             "30%", "≥ 90%"],
            ["BG-06", "Ensure audit-ready records",
             "Audit findings related to HR data", "—", "Zero material findings"],
            ["BG-07", "Strengthen brand experience",
             "Employee NPS for HR tools", "+12", "≥ +45"],
            ["BG-08", "Reduce HR operational cost",
             "HR admin hours / 100 employees / month", "180", "≤ 70"],
            ["BG-09", "Achieve enterprise-grade availability",
             "Platform uptime (rolling 30 days)", "—", "≥ 99.9%"],
            ["BG-10", "Enable data-driven workforce planning",
             "Adoption of analytics dashboards by managers", "—", "≥ 80% MAU"],
        ],
        col_widths=[Inches(0.6), Inches(2.0), Inches(1.7), Inches(1.0), Inches(1.2)]
    )

    add_page_break(doc)

    # 6. Scope
    add_heading(doc, "6. Scope of the Engagement", level=1)
    add_heading(doc, "6.1 In Scope", level=2)
    add_bullets(doc, [
        "Multi-tenant deployment of Perfect Solutions HRMS for the perfectsolutions tenant.",
        "White-label branding: logo, primary/secondary colors, custom domain (hr.perfectsolutions.com), email sender identity, browser tab title.",
        "Modules: Authentication & RBAC, Employee Directory, Time & Attendance, Leave Management, Tickets / Helpdesk, Calendar & Scheduling, Projects & Tasks, Document Vault, Performance Reviews, Reports & Analytics, In-app Chat.",
        "Department configuration: General, Engineering, HR, Sales, Marketing, Finance.",
        "Leave types configured: Vacation, Sick, Personal, Unpaid.",
        "Role-based access control with at minimum: Super Admin, Tenant Admin, HR Manager, People Manager, Employee.",
        "Secure invite-only signup flow tied to tenant-specific admin invite codes.",
        "Data migration from existing spreadsheets for active employees and current leave balances.",
        "End-user training (admins, managers, employees) and operations runbook.",
        "Hypercare period of 30 days post go-live.",
    ])

    add_heading(doc, "6.2 Out of Scope (Phase 1)", level=2)
    add_bullets(doc, [
        "Full payroll calculation and statutory tax filing (handled by external payroll provider; the HRMS will export payroll-ready data).",
        "Recruitment / Applicant Tracking System (ATS) — to be evaluated in Phase 2.",
        "Learning Management System (LMS) integration — Phase 2.",
        "Biometric hardware integration — Phase 2.",
        "Mobile native applications (iOS/Android) — Phase 2 (mobile web is in scope).",
        "Migration of historical (> 24 months) leave and timesheet data.",
    ])

    add_heading(doc, "6.3 Assumptions", level=2)
    add_numbered(doc, [
        "Perfect Solutions will provide a single point of contact (SPOC) empowered to make functional decisions.",
        "Perfect Solutions will provide cleansed master data (employees, departments, managers) prior to migration.",
        "DNS access for hr.perfectsolutions.com will be granted for CNAME verification.",
        "Standard business hours support (9×5) is acceptable post hypercare; 24×7 is a Phase 2 consideration.",
        "All users will access the system using modern evergreen browsers (Chrome, Edge, Safari, Firefox — latest two versions).",
        "English is the only language required at go-live.",
    ])

    add_heading(doc, "6.4 Constraints", level=2)
    add_bullets(doc, [
        "Go-live must occur before the start of the next fiscal year planning cycle.",
        "Data residency: all production data stored in regions compliant with US data protection regulations.",
        "Maximum allowable scheduled downtime is one 2-hour window per month, executed outside business hours.",
        "Implementation budget is fixed; scope changes require formal change-control approval.",
    ])

    add_page_break(doc)

    # 7. Process flows
    add_heading(doc, "7. Business Processes — Current vs Future State", level=1)
    add_para(doc,
        "The table below summarizes the most impactful HR processes that will transition from a manual or semi-digital "
        "current state to a fully automated future state on Perfect Solutions HRMS.")
    add_table(doc,
        ["Process", "Current State (As-Is)", "Future State (To-Be)", "Expected Benefit"],
        [
            ["Employee Onboarding",
             "Email-driven, paper forms, manual asset assignment.",
             "Digital onboarding workflow, pre-arrival document collection, automatic account provisioning.",
             "Onboarding cycle reduced from 5 days to 1 day."],
            ["Time Tracking",
             "Spreadsheets emailed weekly to managers.",
             "In-app daily/weekly entry, manager approvals, project tagging, exception alerts.",
             "Billing accuracy +15%; overtime visibility same-day."],
            ["Leave Requests",
             "Email + spreadsheet ledger maintained by HR.",
             "Self-service request, real-time balance, automated routing, calendar sync.",
             "Cycle time < 1 day; zero balance disputes."],
            ["IT Helpdesk",
             "Shared mailbox; no SLA tracking.",
             "Ticketing module with categories, SLAs, comments, attachments, satisfaction rating.",
             "First response < 2h; CSAT measurable."],
            ["Performance Reviews",
             "Annual, paper-based, Q4-only.",
             "Continuous feedback, quarterly check-ins, OKR tracking.",
             "Higher engagement, retention insight."],
            ["Audit Response",
             "Manual data pull from multiple sources.",
             "Single export with immutable audit log.",
             "Audit prep time reduced ~70%."],
        ],
        col_widths=[Inches(1.2), Inches(1.6), Inches(2.0), Inches(1.6)]
    )

    add_image(doc, f"{ASSETS}/leave_flow.png", width_inches=6.4,
              caption="Figure 3 — Future-state Leave Request Workflow")

    add_screenshot(doc, f"{SCREENSHOTS_DIR}/07_admin_leave.png",
                   "Live Capture — Leave Approval Queue (Manager View)",
                   "Figure 7 — Pending leave requests with policy context for one-click decisions.")
    add_screenshot(doc, f"{SCREENSHOTS_DIR}/08_admin_tickets.png",
                   "Live Capture — Tickets / Internal Helpdesk",
                   "Figure 8 — Categorized ticket queue with SLA visibility.")

    add_page_break(doc)

    # 8. Business Requirements
    add_heading(doc, "8. Business Requirements", level=1)
    add_para(doc,
        "Each business requirement is uniquely identified (BR-NNN), categorized, prioritized using MoSCoW, and traced "
        "to the corresponding business goal. Functional decomposition is provided in the companion FRD "
        "(PS-FRD-2026-001).")

    brs = [
        ("BR-001", "Branding", "The platform must display Perfect Solutions branding (logo, colors, tab title, custom domain) consistently across every page and outbound email.", "Must", "BG-07"),
        ("BR-002", "Access", "Only invited users with a valid tenant-specific admin invite code may sign up as administrators.", "Must", "BG-06"),
        ("BR-003", "Access", "All users must authenticate before accessing any non-public route.", "Must", "BG-06"),
        ("BR-004", "Access", "The system shall enforce role-based access control with at minimum five roles.", "Must", "BG-06"),
        ("BR-005", "Employees", "HR shall be able to create, update, deactivate, and search employee records.", "Must", "BG-01"),
        ("BR-006", "Time", "Employees shall log time daily/weekly with project tagging and submit for manager approval.", "Must", "BG-02"),
        ("BR-007", "Time", "Managers shall approve, reject, or request changes on submitted timesheets.", "Must", "BG-02"),
        ("BR-008", "Leave", "Employees shall view current leave balance and submit leave requests.", "Must", "BG-03"),
        ("BR-009", "Leave", "Managers shall approve/reject leave with automated notifications.", "Must", "BG-03"),
        ("BR-010", "Leave", "The system shall prevent overlapping or balance-exceeding leave requests.", "Must", "BG-03"),
        ("BR-011", "Tickets", "Employees shall raise support tickets categorized by department.", "Must", "BG-04"),
        ("BR-012", "Tickets", "Tickets shall track SLA against configured targets.", "Should", "BG-04"),
        ("BR-013", "Calendar", "Approved leaves and company events shall surface in a unified calendar.", "Must", "BG-05"),
        ("BR-014", "Documents", "Employees shall securely access pay slips, policies, and personal documents.", "Must", "BG-01"),
        ("BR-015", "Reports", "HR and Finance shall export workforce, attendance, and leave reports in CSV/Excel.", "Must", "BG-10"),
        ("BR-016", "Reports", "Leadership shall view real-time dashboards by department, role, and project.", "Should", "BG-10"),
        ("BR-017", "Notifications", "All approval-driven workflows shall produce in-app and email notifications.", "Must", "BG-03"),
        ("BR-018", "Audit", "Every state-changing action shall be recorded in an immutable audit log retained ≥ 7 years.", "Must", "BG-06"),
        ("BR-019", "Security", "Passwords shall be hashed using a modern adaptive algorithm; PII encrypted at rest.", "Must", "BG-06"),
        ("BR-020", "Security", "Failed login attempts shall be rate-limited and lock accounts after 5 consecutive failures.", "Must", "BG-06"),
        ("BR-021", "Availability", "The platform shall maintain ≥ 99.9% monthly uptime.", "Must", "BG-09"),
        ("BR-022", "Performance", "95% of API requests shall return in under 500 ms under normal load.", "Should", "BG-09"),
        ("BR-023", "Usability", "The UI shall be fully responsive and meet WCAG 2.1 AA accessibility.", "Should", "BG-07"),
        ("BR-024", "Self-service", "Employees shall accomplish all routine HR actions without HR intervention.", "Must", "BG-05"),
        ("BR-025", "Integrations", "The platform shall expose REST APIs for payroll, SSO, and BI tools.", "Should", "BG-08"),
        ("BR-026", "Compliance", "The platform shall provide data export and deletion capabilities to satisfy data subject rights.", "Must", "BG-06"),
        ("BR-027", "Localization", "Time zones shall default to America/New_York with per-employee override.", "Should", "BG-07"),
        ("BR-028", "Mobile", "All employee-facing flows shall be usable on mobile browsers (≥ 360 px width).", "Must", "BG-07"),
        ("BR-029", "Tenant Isolation", "No data of any other tenant shall be accessible from the perfectsolutions tenant.", "Must", "BG-06"),
        ("BR-030", "Disaster Recovery", "RPO ≤ 1 hour, RTO ≤ 4 hours.", "Must", "BG-09"),
    ]
    add_table(doc,
        ["ID", "Category", "Requirement", "Priority", "Goal"],
        [list(b) for b in brs],
        col_widths=[Inches(0.6), Inches(0.95), Inches(3.2), Inches(0.7), Inches(0.6)]
    )

    add_page_break(doc)

    # 9. Risks
    add_heading(doc, "9. Risk Analysis & Mitigation", level=1)
    add_para(doc,
        "Risks are scored on Likelihood (L) and Impact (I) on a 1–5 scale. Risk Score = L × I. Risks scoring ≥ 12 are "
        "tracked weekly by the steering committee.")
    add_table(doc, ["ID", "Risk", "L", "I", "Score", "Mitigation", "Owner"],
        [
            ["R-01", "Resistance to change from long-tenured employees", 3, 4, 12,
             "Change-management plan, executive sponsorship, hands-on training.", "HR Lead"],
            ["R-02", "Master-data quality issues delay migration", 4, 4, 16,
             "Data-cleansing sprint and phased migration with rollback.", "PMO"],
            ["R-03", "DNS / CNAME setup delays custom domain go-live", 2, 3, 6,
             "Pre-validate DNS access early; fallback to *.hr.perfectsolutions.com.", "IT"],
            ["R-04", "Security incident exposes PII", 2, 5, 10,
             "Encryption at rest + in transit, RBAC, audit log, SOC monitoring.", "InfoSec"],
            ["R-05", "Vendor LLM costs exceed budget", 3, 2, 6,
             "Universal Key with budget alerts; cap monthly spend per feature.", "Finance"],
            ["R-06", "Browser compatibility regressions", 3, 2, 6,
             "Automated cross-browser test suite in CI.", "Engineering"],
            ["R-07", "Insufficient mobile experience reduces adoption", 3, 4, 12,
             "Mobile-first design review for every employee flow.", "Product"],
            ["R-08", "Unplanned downtime during peak hours", 2, 5, 10,
             "Blue-green deploys, off-hours maintenance windows, SRE on-call.", "DevOps"],
            ["R-09", "Cross-tenant data leakage", 1, 5, 5,
             "Tenant-scoped queries, automated isolation tests in CI.", "Engineering"],
            ["R-10", "Audit failure due to incomplete logging", 2, 5, 10,
             "Immutable append-only audit log with periodic verification.", "Compliance"],
        ],
        col_widths=[Inches(0.5), Inches(2.0), Inches(0.4), Inches(0.4), Inches(0.6), Inches(2.2), Inches(0.6)]
    )

    add_page_break(doc)

    # 10. Compliance
    add_heading(doc, "10. Regulatory & Compliance Considerations", level=1)
    add_bullets(doc, [
        "U.S. labor recordkeeping (Fair Labor Standards Act): time records retained for ≥ 3 years; payroll-ready exports.",
        "U.S. state-specific paid-sick-leave rules: leave types configurable per policy.",
        "Data privacy: explicit consent capture, right to access and delete, breach notification within 72 hours.",
        "SOC 2 Type II posture: annual third-party assessment of security, availability, confidentiality.",
        "Internal IT policies: SSO-readiness, password complexity, session timeout (≤ 30 minutes idle).",
    ])

    # 11. Glossary
    add_heading(doc, "11. Glossary", level=1)
    add_table(doc, ["Term", "Definition"],
        [
            ["BRD", "Business Requirements Document"],
            ["FRD", "Functional Requirements Document"],
            ["HRMS", "Human Resource Management System"],
            ["MoSCoW", "Prioritization framework: Must / Should / Could / Won't"],
            ["RBAC", "Role-Based Access Control"],
            ["RPO", "Recovery Point Objective — max acceptable data loss measured in time"],
            ["RTO", "Recovery Time Objective — max acceptable downtime to restore service"],
            ["SLA", "Service Level Agreement"],
            ["SPOC", "Single Point Of Contact"],
            ["SaaS", "Software as a Service"],
            ["Tenant", "Isolated logical instance within the multi-tenant SaaS platform"],
            ["UAT", "User Acceptance Testing"],
            ["WCAG", "Web Content Accessibility Guidelines"],
        ],
        col_widths=[Inches(1.4), Inches(5.0)]
    )

    add_page_break(doc)

    # 12. Sign-off
    add_heading(doc, "12. Approval & Sign-off", level=1)
    add_para(doc,
        "The undersigned acknowledge that this Business Requirements Document represents the agreed-upon business "
        "needs of Perfect Solutions for the Perfect Solutions HRMS implementation. Material changes to this document shall "
        "be processed via formal change control.")
    add_para(doc, "")
    add_table(doc, ["Role", "Name", "Signature", "Date"],
        [
            ["Executive Sponsor", "Director, IT Services — Perfect Solutions", "_____________________", "__________"],
            ["Business Owner", "Head of Human Resources — Perfect Solutions", "_____________________", "__________"],
            ["Financial Approver", "Chief Financial Officer — Perfect Solutions", "_____________________", "__________"],
            ["Implementation Lead", "IT Delivery", "_____________________", "__________"],
            ["Product Manager", "Product Office", "_____________________", "__________"],
        ],
        col_widths=[Inches(1.5), Inches(2.6), Inches(1.6), Inches(0.8)]
    )

    out = f"{OUT_DIR}/Perfect_Solutions_BRD_v1.0.docx"
    add_portal_appendix(doc)
    doc.save(out)
    print(f"BRD DOCX saved → {out}")
    return out


if __name__ == "__main__":
    build_brd()
